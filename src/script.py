#!/usr/bin/python3
import os
import cfscrape
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches

#Instancia o objeto responsável por "quebrar" o cloudfare do URI
scraper = cfscrape.create_scraper()

#Instancia o objeto para escrever num arquivo .docx
document = Document()
document.add_heading('URI - Questões', 0)

with open('../input/lista.txt') as file:
	print('Running...')
	for line in file:
		request = scraper.get('https://www.urionlinejudge.com.br/repository/UOJ_'+line.strip()+'.html')

		#Caso encontre algum erro 404, ele passa para o próximo elemento do arquivo
		if request.status_code==404: continue
		request.encoding="utf-8"
		soup = BeautifulSoup(request.text, 'html.parser')
		description_ = soup.find(class_='description')
		input_ = soup.find(class_='input')
		output_ = soup.find(class_='output')

		document.add_heading('Questão '+line.strip(), level=1)

		document.add_heading('Descrição', level=2)
		document.add_paragraph(description_.text.strip())

		document.add_heading('Entrada', level=3)
		document.add_paragraph(input_.text.strip())

		document.add_heading('Saída', level=4)
		document.add_paragraph(output_.text.strip())

print('The program has been successfully executed!')
#Cria o diretório de saída do arquivo, caso não exista;
if not os.path.exists(os.path.join('..','output')): os.mkdir(os.path.join('..','output'))

document.save(os.path.join('..','output','questions.docx'))